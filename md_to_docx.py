"""Converts a Markdown file to a Word document.

Iterates through the lines of the Markdown file and parses them, adding headings, paragraphs, and formatting based on the Markdown syntax. Headings, bold text, empty lines, etc. are all rendered appropriately in the Word document.

Args:
    file_path (str): The path to the Markdown file to convert.
"""
from docx import Document
from docx.shared import Pt
import re

# Add a heading to the document.
def add_heading(document, text, level):
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(24)
    elif level == 2:
        run.font.size = Pt(20)
    else:
        run.font.size = Pt(16)

# Add a paragraph to the document.
def add_paragraph(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)

# Add a hyperlink to the document.
def parse_markdown_line(document, line):
    if line.startswith('## '):
        add_heading(document, line[3:], level=1)
    elif line.startswith('### '):
        add_heading(document, line[4:], level=2)
    elif line.startswith('- **'):
        match = re.match(r'- \*\*(.*):\*\* (.*)', line)
        if match:
            add_paragraph(document, f'{match.group(1)}: {match.group(2)}')
    elif line.strip() == '':
        document.add_paragraph()  # Add an empty paragraph for empty lines, or you could just pass
    else:
        add_paragraph(document, line)

# Convert the Markdown file to a Word document.
def markdown_to_word(file_path):
    document = Document()
    
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            parse_markdown_line(document, line)
    
    # You can modify the output file name or make it parameterized as well
    output_file_name = file_path.replace('.md', '.docx')
    document.save(output_file_name)
    print(f'Document saved as {output_file_name}')

# Example usage
file_path = 'path/ to /file.md'  # Replace with the actual Markdown file path
markdown_to_word(file_path)
