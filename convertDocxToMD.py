# A try to avoide Pandoc and use python-docx instead
from docx import Document
from docx.oxml.ns import qn

def convert_docx_to_markdown(docx_file, md_file):
    doc = Document(docx_file)
    md_content = []

    def get_heading_level(style):
        if style.startswith('Heading'):
            try:
                level = int(style.replace('Heading', '').strip())
                return level
            except ValueError:
                return None
        return None

    def process_paragraph(paragraph):
        style = paragraph.style.name
        level = get_heading_level(style)
        if level:
            md_content.append(f"{'#' * level} {paragraph.text}\n")
        else:
            md_content.append(f"{paragraph.text}\n")

    def process_table(table):
        for i, row in enumerate(table.rows):
            row_content = []
            for cell in row.cells:
                row_content.append(cell.text.strip().replace('\n', ' '))
            md_content.append(' | '.join(row_content) + '\n')
            if i == 0:
                md_content.append(' | '.join(['---'] * len(row.cells)) + '\n')

    def iter_block_items(parent):
        """Yield each paragraph and table child within *parent*, in document order."""
        for child in parent.element.body:
            if child.tag == qn('w:p'):
                yield 'paragraph', doc.paragraphs[next(paragraph_count)]
            elif child.tag == qn('w:tbl'):
                yield 'table', doc.tables[next(table_count)]

    paragraph_count = iter(range(len(doc.paragraphs)))
    table_count = iter(range(len(doc.tables)))

    for block_type, block in iter_block_items(doc):
        if block_type == 'paragraph':
            process_paragraph(block)
        elif block_type == 'table':
            process_table(block)

    with open(md_file, 'w', encoding='utf-8') as f:
        f.writelines(md_content)

# Usage example
convert_docx_to_markdown('data/example.docx', 'tmp/example.md')
